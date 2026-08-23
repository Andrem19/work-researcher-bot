import os, sys
from docx import Document
from docx.shared import Pt, Mm

path = r'D:\PYTHON\WORK_RESEARCHER_MCP\CV_collection\Cover_Letter_Geophysical_Logging.docx'

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

# Header
p = doc.add_paragraph()
run = p.add_run('Andrew Remniow')
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run('Blackpool, UK | +44 7838 228012 | 7255591@gmail.com | Full UK Driving Licence & Own Vehicle')
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('23 August 2026')
doc.add_paragraph('Ernest Gordon Recruitment')

p = doc.add_paragraph()
run = p.add_run('Re: Graduate/Trainee Field Service Engineer (Geophysical Logging) — Llandudno (LL30)')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph('Dear Hiring Manager,')

doc.add_paragraph(
    'I am writing to express my strong interest in the Graduate/Trainee Field Service Engineer position in Geophysical Logging. '
    'With a BSc in Geology (UK NARIC: RQF Level 6) and hands-on field experience in engineering geology, I am confident that my '
    'background aligns well with the requirements of this role and that I can make a meaningful contribution to your team from day one.'
)

p = doc.add_paragraph()
run = p.add_run('Relevant Field Experience')
run.bold = True
doc.add_paragraph(
    'My field experience includes borehole supervision, trial pits, soil and rock sampling, core and cuttings handling, '
    'groundwater observations and basic lithological logging. I have supported geotechnical laboratory operations and '
    'maintained accurate field records, daily logs and factual summaries throughout my work. This directly relates to the '
    'geophysical logging work your team undertakes, and I am familiar with the disciplined, safety-focused approach '
    'required for borehole and wellsite operations. I am accustomed to working in variable outdoor conditions and '
    'maintaining high standards of data quality and operational safety at all times.'
)

p = doc.add_paragraph()
run = p.add_run('Readiness for Field-Based and Travel-Intensive Work')
run.bold = True
doc.add_paragraph(
    'I hold a full UK driving licence and my own vehicle, and I am fully prepared and enthusiastic about field-based, '
    'travel-intensive work. I am comfortable with rotational work patterns, onshore assignments and extended periods '
    'away from home. My HSE awareness and safety mindset are central to how I operate in the field, and I am confident '
    'working in physically demanding and variable environments. I understand that geophysical logging often requires '
    'working at remote wellsite locations, sometimes on extended rotations, and I am fully committed to this type of '
    'work schedule. I am reliable, punctual and able to maintain focus and accuracy during long field shifts.'
)

p = doc.add_paragraph()
run = p.add_run('Additional Technical Skills')
run.bold = True
doc.add_paragraph(
    'In addition to my geoscience background, I bring strong digital skills including Python scripting, SQL and PostgreSQL, '
    'and advanced Excel. These support data handling, report generation and QA/QC processes, complementing the '
    'technical and analytical demands of geophysical data collection and interpretation. I am comfortable working with '
    'structured data, producing clear and accurate documentation, and automating repetitive reporting tasks. My '
    'experience with structured record-keeping and QA/QC from both geological field operations and software '
    'development gives me a strong foundation for the data-intensive aspects of geophysical logging.'
)

p = doc.add_paragraph()
run = p.add_run('Motivation and Commitment')
run.bold = True
doc.add_paragraph(
    'As a trainee candidate, I am eager to undergo the extensive training you offer and to develop into a skilled '
    'field service engineer. I learn quickly, adapt well to new environments and equipment, and am motivated to build '
    'a long-term career in geophysical logging. I am a British citizen with an unrestricted right to work in the UK '
    'and am available to start at short notice. I am based in Blackpool but am fully mobile and willing to travel to '
    'Llandudno or any other base location as required for field operations.'
)

doc.add_paragraph(
    'I would welcome the opportunity to discuss how my background, field experience and enthusiasm can contribute to '
    'your team. I am available for an interview at your convenience and can be reached on +44 7838 228012 or by email '
    'at 7255591@gmail.com. Thank you for considering my application. I look forward to hearing from you.'
)

doc.add_paragraph()
doc.add_paragraph('Yours sincerely,')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Andrew Remniow')
run.bold = True

doc.save(path)
sys.stdout.write(f'exists: {os.path.exists(path)} size: {os.path.getsize(path)}\n')
sys.stdout.flush()
