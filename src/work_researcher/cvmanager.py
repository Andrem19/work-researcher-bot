"""CV collection manager: active profile's local CV directory + SQLite index."""

from __future__ import annotations

import asyncio
import hashlib
import re
from pathlib import Path

import aiosqlite

from .config import Settings
from .textutils import clean

CV_EXTENSIONS = {".docx", ".pdf", ".doc"}

DOMAIN_TAGS = {
    "data_analytics": [
        "sql",
        "tableau",
        "power bi",
        "powerbi",
        "python",
        "pandas",
        "excel",
        "dashboard",
        "analytics",
        "data analyst",
        "data analysis",
        "etl",
        "visualization",
        "visualisation",
        "machine learning",
        "statistics",
        "kpi",
        "reporting",
    ],
    "geology": [
        "geolog",
        "geotech",
        "geoscience",
        "site investigation",
        "gis",
        "drilling",
        "borehole",
        "strata",
        "mineral",
        "exploration",
        "logging",
        "contamination",
        "hydrogeolog",
        "rock",
        "soil",
        "fieldwork",
        "mapping",
        "engineering geologist",
        "quarry",
        "aggregates",
    ],
    "engineering": [
        "engineer",
        "cad",
        "autocad",
        "civil",
        "structural",
        "mechanical",
        "project management",
        "design",
    ],
}

_LANG_HINTS = {
    "en": ("education", "experience", "skills", "summary", "references"),
    "ru": ("образование", "опыт", "навыки", "резюме", "о себе"),
}


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return _extract_docx(path)
    if path.suffix.lower() == ".pdf":
        return _extract_pdf(path)
    return ""


def classify(text: str) -> list[str]:
    low = text.lower()
    tags = []
    for domain, keywords in DOMAIN_TAGS.items():
        hits = sum(1 for kw in keywords if kw in low)
        if hits >= 3 or (hits >= 2 and domain in ("data_analytics", "geology")):
            tags.append(domain)
    return tags


def detect_language(text: str) -> str | None:
    low = text.lower()
    best, best_hits = None, 0
    for lang, hints in _LANG_HINTS.items():
        hits = sum(1 for h in hints if h in low)
        if hits > best_hits:
            best, best_hits = lang, hits
    return best


_NAME_RE = re.compile(r"^[A-Z][a-zA-Z.'-]+(?:\s+[A-Z][a-zA-Z.'-]+){1,3}$")


def guess_name(text: str) -> str | None:
    """The first standalone line that looks like a person's full name."""
    for line in text.splitlines()[:12]:
        line = clean(line)
        if 5 <= len(line) <= 45 and _NAME_RE.match(line):
            if not re.search(r"(cv|curriculum|resume|address|phone|email|profile)", line, re.I):
                return line
    return None


def scan_local(settings: Settings) -> list[Path]:
    if not settings.cv_dir.exists():
        return []
    return sorted(
        p
        for p in settings.cv_dir.iterdir()
        if p.is_file() and p.suffix.lower() in CV_EXTENSIONS and not p.name.startswith("~$")
    )


async def index_cvs(settings: Settings, force: bool = False, prune_missing: bool = True) -> dict:
    """Make the CV index mirror the active profile's local directory."""
    from . import persistence as db

    files = scan_local(settings)
    scanned, updated, skipped, failed, removed = 0, 0, 0, 0, 0
    async with db.connect(settings.db_path) as conn:
        if prune_missing:
            local_paths = {path.resolve() for path in files}
            cursor = await conn.execute("SELECT id, path FROM cvs")
            missing_ids = []
            for row in await cursor.fetchall():
                indexed_path = Path(row["path"])
                try:
                    same_folder = indexed_path.parent.resolve() == settings.cv_dir.resolve()
                    is_missing = indexed_path.resolve() not in local_paths
                except OSError:
                    same_folder = indexed_path.parent == settings.cv_dir
                    is_missing = indexed_path not in files
                if same_folder and is_missing:
                    missing_ids.append((row["id"],))
            if missing_ids:
                await conn.executemany("DELETE FROM cvs WHERE id=?", missing_ids)
                removed = len(missing_ids)
        for path in files:
            scanned += 1
            stat = path.stat()
            key = f"{stat.st_size}:{int(stat.st_mtime)}"
            cur = await conn.execute("SELECT mtime FROM cvs WHERE path=?", (str(path),))
            row = await cur.fetchone()
            if row and row["mtime"] == key and not force:
                skipped += 1
                continue
            try:
                text = await asyncio.to_thread(extract_text, path)
            except Exception as exc:
                failed += 1
                await db.upsert_cv(
                    conn,
                    {
                        "filename": path.name,
                        "path": str(path),
                        "sha256": None,
                        "size": stat.st_size,
                        "mtime": key,
                        "text_preview": f"[parse failed: {exc}]",
                        "full_text": "",
                        "tags": "[]",
                        "language": None,
                        "name_guess": None,
                    },
                )
                continue
            updated += 1
            await db.upsert_cv(
                conn,
                {
                    "filename": path.name,
                    "path": str(path),
                    "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                    "size": stat.st_size,
                    "mtime": key,
                    "text_preview": clean(text)[:600],
                    "full_text": text[:20000],
                    "tags": str(classify(text)),
                    "language": detect_language(text),
                    "name_guess": guess_name(text),
                },
            )
        await conn.commit()
    return {
        "files_found": scanned,
        "indexed": updated,
        "unchanged": skipped,
        "failed": failed,
        "removed_missing": removed,
        "cv_dir": str(settings.cv_dir),
    }


JOB_DOMAIN_HINTS = {
    "geology": (
        "geolog",
        "geoscience",
        "geotech",
        "geophysic",
        "hydrogeolog",
        "borehole",
        "drilling",
        "mud log",
        "site investigation",
        "contaminated land",
        "mineral",
        "mining",
        "quarry",
        "fieldwork",
        "geophysical logging",
        "geo-environmental",
        "ground investigation",
    ),
    "data_analytics": (
        "data analyst",
        "analytics",
        "sql",
        "tableau",
        "power bi",
        "dashboard",
        "insight",
        "reporting analyst",
        "business intelligence",
        "data scientist",
    ),
}


def job_domains(job_text: str) -> set[str]:
    low = (job_text or "").lower()
    return {d for d, words in JOB_DOMAIN_HINTS.items() if any(w in low for w in words)}


async def recommend_cv(conn: aiosqlite.Connection, job: dict, limit: int = 3) -> list[dict]:
    """Rank CVs for a job: domain-tag match FIRST (a geology role must not
    surface the Data Analyst CV just because it mentions more keywords),
    then keyword coverage."""
    from .textutils import query_terms, term_coverage

    cur = await conn.execute("SELECT id, filename, tags, full_text, indexed_at FROM cvs")
    rows = await cur.fetchall()
    job_text = " ".join(
        str(job.get(k) or "") for k in ("title", "description", "company", "location_text")
    )
    domains = job_domains(job_text)
    scored = []
    for r in rows:
        tags = re.findall(r"[a-z_]+", r["tags"] or "")
        tag_bonus = 0.0
        for domain in domains:
            if domain in tags:
                tag_bonus += 0.35  # CV built for this domain
            elif tags and domain in JOB_DOMAIN_HINTS:
                tag_bonus -= 0.15  # CV from a different domain
        cov = term_coverage(query_terms(job_text)[:15], r["full_text"] or "")
        score = round(min(1.0, cov + tag_bonus), 2)
        scored.append(
            {
                "cv_id": r["id"],
                "filename": r["filename"],
                "tags": tags,
                "score": score,
                "indexed_at": r["indexed_at"],
                "domain_match": sorted(domains & set(tags)),
            }
        )
    scored.sort(key=lambda x: -x["score"])
    # tie-break: a CV FILENAME named after the job domain (e.g.
    # Engineering_Geology.docx for a geology role) outranks generic files
    domain_words = [w.replace("_", " ") for d in domains for w in (d, d + " cv")]
    if domain_words:

        def _named(rec: dict) -> int:
            fn = rec["filename"].lower().replace("-", " ")
            return 0 if any(w in fn for w in domain_words) else 1

        scored.sort(key=lambda x: (-x["score"], _named(x)))
    return scored[:limit]
