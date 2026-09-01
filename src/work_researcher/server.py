"""MCP server wiring: search, CVs, applications and the embedded browser.

Tool surface is deliberately COMPACT (24 tools, prefix-grouped) so mid-size
agents can hold it in mind:

  search     get_status, search_jobs, get_job, submit_job_observations
  cv         list_cvs, sync_cvs
  apply      start_application, record_application, list_applications,
             check_applied
  browser    browser_open, browser_snapshot, browser_form, browser_click,
             browser_set, browser_type, browser_upload, browser_press,
             browser_wait, browser_screenshot, browser_eval, browser_tabs,
             browser_close

Stdio transport; logs go to stderr only (stdout carries the protocol).
"""

from __future__ import annotations

import asyncio
import json
import logging
import sys
from typing import Any, Literal

from mcp.server import MCPServer

from . import dedup
from . import persistence as db
from . import tracker as tracker_mod
from .config import Settings, ensure_dirs, load_settings
from .domain import JobCard, SearchParams
from .providers import BROWSER_ONLY_NOTES, run_search
from .ranking import score_job
from .textutils import job_hash

INSTRUCTIONS = (
    "Work Researcher MCP: UK job search + application engine (29 compact tools). "
    "SINGLE CANDIDATE: this server is configured only for Andrey Remnev. "
    "ADAPTIVE SEARCH RESPONSES: on search_jobs and get_job pass context_window "
    "equal to your model's advertised context size. Local Qwen 3.8 27B MUST pass "
    "context_window=78000 (or response_profile='compact'); models around 100k-256k "
    "use balanced; models with 300k+ or 1M context use wide. The server selects a "
    "compact/balanced/wide page automatically, stores that policy with search_id, "
    "and returns next_page arguments. If model context is unknown, omit it for the "
    "safe balanced default. Do not launch more than 2 fresh searches in one tool "
    "round or fetch every page without first ranking the current page. get_job is "
    "compact by default; request full descriptions only for finalists. No search "
    "data is discarded by response pagination. "
    "USER WORKFLOW: (1) search_jobs(query or profile) — e.g. 'Data Analytics' or "
    "'Field Geologist Engineer'; (2) present the ranked list to the user — every "
    "row MUST include the posted_by column (agency vs direct employer) AND a "
    "short description (2-3 sentences: what the job involves, key duties); if "
    "the description field is null, call fetch_job_description for the top "
    "picks BEFORE presenting; (3) the "
    "user picks "
    "vacancies; (4) start_application(job_id) per pick — it refuses double "
    "applications (memory across sessions) and returns URL + method + CV + "
    "applicant profile + site playbook; (5) ensure login with browser_login(url) "
    "— 'Continue with Google' picking the pre-approved account from config "
    "[auth] WITHOUT asking the user (2FA/captcha → stop and ask); (6) drive the "
    "submission with browser_* tools; (7) record_application(status='submitted', "
    "evidence={screenshot}). LOCATION INTELLIGENCE: results carry work_mode / "
    "distance_miles / location_status (home comes from the active profile). Remote jobs "
    "are searched UK-wide; on-site jobs outside max_commute_miles "
    "are flagged mismatch — never submit those without explicit user approval. "
    "LOCATION IS WORK-MODE-AWARE: on_site (daily office) must be within "
    "daily_commute_miles (default 25); hybrid/field/unknown within "
    "occasional_commute_miles (default 50); remote = any distance. By default "
    "drop_mismatch=true removes on-site-mismatch jobs from the results "
    "(location_skipped counts them) — pass drop_mismatch=false to see them. "
    "BLOCKLIST: when the user says 'never apply to X' → manage_blocklist "
    "action=add kind=company; blocked employers are hidden from results and "
    "start_application refuses them. TRAINING-AD GUARD: paid course ads (where "
    "the candidate pays: 'training course', Netcom-style providers, fee/loan "
    "language, training company + trainee title + no salary) are EXCLUDED "
    "automatically — training_offers_skipped counts them. When the user asks "
    "for 'junior/trainee with training', show only real paid vacancies "
    "(salary-present apprenticeships count); never propose paying for courses. "
    "TOKEN DISCIPLINE: never re-run the same search_jobs query twice — the tool returns note_duplicate with the existing search_id; PAGE it via search_id+offset instead. Never query the SQLite DB via Bash — use list_stored_jobs (filter stored jobs by company/title/location/source). Batch fetch_job_description(job_ids=[...]) up to 10 at once instead of one-by-one calls. REQUIREMENTS CHECK: jobs with hard requirements the user does NOT meet "
    "(from CVs) are flagged requirements_status=gap and dropped by default "
    "(req_skipped counts them; drop_req_gap=false to see). Never propose a job "
    "requiring e.g. AAT Level 2 if the user's CVs lack it — check "
    "requirements_unmet before listing. "
    "BROWSER SPEED PROTOCOL: browser_open "
    "returns the first snapshot; every click/set/type/upload RETURNS a fresh "
    "snapshot — never re-snapshot between steps; use the element numbers from "
    "the last result. On application forms call browser_form once, then "
    "browser_set by field number. browser_snapshot(filter_text='Apply') finds "
    "buttons; text_chars=6000 reads descriptions. UPLOADS: browser_upload sets "
    "files DIRECTLY on input[type=file] (hidden inputs included — Totaljobs "
    "style) and falls back to the native chooser; cover letters: "
    "make_cover_letter(text) → DOCX >8KB (board minimum) → browser_upload. "
    "WIZARD PROTOCOL (Reed apply questions and any modal-based form): use browser_snapshot(modal_only=true) — it isolates the ACTIVE wizard (question text + numbered controls, hidden templates like 'Session expired' excluded) — then browser_set/browser_click by number. NEVER hand-roll modal-finding JS via browser_eval (slow, fragile, and hidden templates mislead it). NEVER close a wizard modal: its X loses ALL progress and it restarts from Q1. NEVER click unnamed/empty buttons. Required multi-selects without a visible none-option: SCROLL the list first ('prefer not to say' / 'none of these' usually sits below the fold). CRITICAL: NEVER use the separate playwright MCP (mcp__playwright__*) for job pages or applications — its browser has NO logins and is a DIFFERENT context; always use this server's browser_* tools exclusively. IMPORTANT: the application form lives in THIS server's browser — never "
    "drive it with a separate Playwright/browser MCP (different context, not "
    "logged in). Only "
    "Indeed/LinkedIn/CV-Library/Glassdoor need the harness's own browser — feed "
    "findings back via submit_job_observations. GOV.UK Work Hub "
    "(jobs.service.gov.uk/jobs) is also browser-only: browser_open "
    "→ fill keywords input (id=keywordsInput) → Search → run the "
    "SCRAPE_LINKS_JS from providers/govuk_workhub.py via browser_eval → "
    "submit_job_observations. Its REMOTE/ONSITE/HYBRID/FIELD_BASED filters "
    "map directly to work_mode. Dedup merges the same vacancy "
    "across boards — check sources[] and already_applied before applying."
)

logger = logging.getLogger("work_researcher")


ResponseProfile = Literal["auto", "compact", "balanced", "wide"]

_SEARCH_RESPONSE_PROFILES = {
    "compact": {"default_limit": 4, "max_limit": 8},
    "balanced": {"default_limit": 12, "max_limit": 20},
    "wide": {"default_limit": 30, "max_limit": 50},
}


def _resolve_response_policy(
    response_profile: ResponseProfile = "auto",
    context_window: int | None = None,
    limit: int | None = None,
) -> dict[str, Any]:
    """Resolve a bounded response page without pretending MCP knows the LLM.

    MCP tool calls do not include the caller model or its context window.  The
    caller can supply either value explicitly; otherwise balanced is the least
    surprising cross-model default.  An explicit limit with no other signal is
    treated as intentional and is used to infer the closest capacity profile.
    """
    if context_window is not None and context_window < 1:
        raise ValueError("context_window must be a positive token count")
    if limit is not None and limit < 1:
        raise ValueError("limit must be at least 1")

    selected = response_profile
    reason = "explicit_profile"
    if selected == "auto":
        if context_window is not None:
            if context_window <= 80_000:
                selected = "compact"
            elif context_window < 300_000:
                selected = "balanced"
            else:
                selected = "wide"
            reason = "context_window"
        elif limit is not None:
            if limit <= _SEARCH_RESPONSE_PROFILES["compact"]["max_limit"]:
                selected = "compact"
            elif limit <= _SEARCH_RESPONSE_PROFILES["balanced"]["max_limit"]:
                selected = "balanced"
            else:
                selected = "wide"
            reason = "explicit_limit"
        else:
            selected = "balanced"
            reason = "safe_default"

    spec = _SEARCH_RESPONSE_PROFILES[selected]
    requested = spec["default_limit"] if limit is None else int(limit)
    page_limit = max(1, min(requested, spec["max_limit"]))
    return {
        "requested_profile": response_profile,
        "profile": selected,
        "reason": reason,
        "context_window": context_window,
        "requested_limit": limit,
        "page_limit": page_limit,
        "max_limit": spec["max_limit"],
    }


def configure_logging(settings: Settings) -> None:
    handler = logging.StreamHandler(sys.stderr)
    handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(name)s: %(message)s"))
    logger.handlers.clear()
    logger.addHandler(handler)
    logger.setLevel(settings.log_level)
    logger.propagate = False
    logging.getLogger("httpx").setLevel(logging.WARNING)
    logging.getLogger("httpcore").setLevel(logging.WARNING)


def create_server(settings: Settings | None = None) -> tuple[MCPServer, Settings]:
    settings = settings or load_settings()
    ensure_dirs(settings)
    configure_logging(settings)
    mcp = MCPServer("work-researcher", instructions=INSTRUCTIONS, version="0.1.0")
    _register_tools(mcp, settings)
    return mcp, settings


def _register_tools(mcp: MCPServer, settings: Settings) -> None:
    from .browser import get_session
    from .cvmanager import index_cvs as _index_cvs
    from .cvmanager import recommend_cv as _recommend_cv

    # ------------------------------------------------------------ search ----
    @mcp.tool()
    async def get_status() -> dict:
        """Health/config snapshot: active candidate, local CV directory, DB
        stats, providers, API keys and saved searches. Call first in a session."""
        async with db.connect(settings.db_path) as conn:
            jobs = await db.count_rows(conn, "jobs")
            apps = await db.count_rows(conn, "applications")
            cvs = await db.count_rows(conn, "cvs")
            searches = await db.count_rows(conn, "searches")
        providers = {}
        for name in ("totaljobs", "reed", "adzuna", "jooble", "earthworks", "findajob"):
            key_ok = True
            if name == "adzuna":
                key_ok = bool(settings.secret("adzuna", "app_id")
                              and settings.secret("adzuna", "app_key"))
            elif name in ("reed", "jooble"):
                key_ok = bool(settings.secret(name, "api_key"))
            providers[name] = {
                "enabled": settings.provider_enabled(name),
                "credentials": key_ok,
                "note": "" if key_ok or name in ("totaljobs", "earthworks", "findajob")
                else "add key in config.toml (SETUP.md); Reed falls back to HTML",
            }
        try:
            import playwright  # noqa: F401

            pw = "installed"
        except ImportError:
            pw = "MISSING"
        return {
            "candidate": {
                "display_name": settings.candidate_name,
                "cv_dir": str(settings.cv_dir),
                "data_dir": str(settings.data_dir),
            },
            "database": {"jobs": jobs, "applications": apps, "cvs": cvs,
                         "searches": searches, "path": str(settings.db_path)},
            "providers": providers,
            "browser_only_sources": BROWSER_ONLY_NOTES,
            "cv_storage": {
                "mode": "local_manual",
                "directory": str(settings.cv_dir),
                "instruction": "Copy CV files here, then call sync_cvs to index them",
                "extensions": [".docx", ".pdf", ".doc"],
            },
            "playwright": pw,
            "browser_profile": str(settings.browser_profile_dir),
            "search_profiles": {k: v.get("query")
                                for k, v in settings.search_profiles.items()},
            "response_sizing": {
                "model_context_not_visible_to_mcp": True,
                "pass_to_search_jobs_and_get_job": "context_window",
                "local_qwen_3_8_27b": 78000,
                "profiles": _SEARCH_RESPONSE_PROFILES,
                "unknown_context_default": "balanced",
            },
            "home": {
                "location": settings.applicant.get("home_location"),
                "max_commute_miles": settings.applicant.get("max_commute_miles"),
                "willing_to_relocate": settings.applicant.get("willing_to_relocate"),
            },
            "applicant_configured": bool(settings.applicant.get("full_name")),
        }

    async def _run_search(params: SearchParams, response_policy: dict[str, Any]) -> dict:
        # duplicate-search guard: the same query re-run within 10 minutes
        # returns the EXISTING search_id instead of hitting the boards again
        # (weak models repeat identical searches; this saves tokens and time)
        async def _recent_duplicate() -> str | None:
            from datetime import UTC, datetime, timedelta

            try:
                async with db.connect(settings.db_path) as aconn:
                    cutoff = (datetime.now(tz=UTC) - timedelta(minutes=10)) \
                        .isoformat(timespec="seconds")
                    cur = await aconn.execute(
                        "SELECT id, params FROM searches "
                        "WHERE created_at >= ? ORDER BY created_at DESC LIMIT 8",
                        (cutoff,))
                    for row in await cur.fetchall():
                        try:
                            p = json.loads(row["params"])
                        except (TypeError, ValueError):
                            continue
                        if (p.get("query", "").lower() == params.query.lower()
                                and (p.get("location") or "")
                                == (params.location or "")):
                            return row["id"]
            except Exception:  # noqa: BLE001 - the guard must never break search
                return None
            return None

        _dup_id = await _recent_duplicate()
        if _dup_id:
            out = await _page_results(_dup_id, response_policy, offset=0)
            out["note_duplicate"] = (
                "identical query ran in the last 10 minutes — returning the "
                "existing search; PAGE it with search_id + offset instead of "
                "re-searching")
            return out
        from . import geo as geo_mod

        async with db.connect(settings.db_path) as conn:
            cards_by_provider, reports = await run_search(settings, params.model_dump())
            all_cards = [c for cards in cards_by_provider.values() for c in cards]
            pool = await dedup.load_pool(conn)
            resolution, merged = dedup.resolution_map(all_cards, pool)

            home = await geo_mod.home_geo(settings, conn)
            home_settings = dict(
                home_lat=home and home["lat"], home_lon=home and home["lon"],
                home_location=settings.applicant.get("home_location", "home"),
                max_commute_miles=int(settings.applicant.get("max_commute_miles", 40)),
                daily_commute_miles=int(settings.applicant.get(
                    "daily_commute_miles", 25)),
                occasional_commute_miles=int(settings.applicant.get(
                    "occasional_commute_miles", 50)),
                willing_to_relocate=bool(settings.applicant.get("willing_to_relocate")),
                relocate_areas=list(settings.applicant.get("relocate_areas") or []),
                location_policy=params.location_policy,
            )
            rep_cards: dict[str, JobCard] = {}
            for card in all_cards:
                ch = job_hash(card.title, card.company, card.location_text, card.salary_min)
                canonical = resolution.get(ch, ch)
                cur = rep_cards.get(canonical)
                if cur is None or len(card.description or "") > len(cur.description or ""):
                    rep_cards[canonical] = card

            async def _eval(card: JobCard) -> dict:
                work_mode = geo_mod.classify_work_mode(card.title, card.description)
                lat, lon = card.extra.get("latitude"), card.extra.get("longitude")
                if lat is None:
                    g = await geo_mod.geocode(settings, card.location_text, conn)
                    if g:
                        lat, lon = g["lat"], g["lon"]
                return geo_mod.evaluate_location(
                    work_mode=work_mode, job_lat=lat, job_lon=lon,
                    job_location=card.location_text, **home_settings)

            sem = asyncio.Semaphore(5)

            async def _eval_guarded(card):
                async with sem:
                    try:
                        return await _eval(card)
                    except Exception:  # noqa: BLE001 - geo must never break search
                        return {"work_mode": None, "distance_miles": None,
                                "location_status": "unknown", "reason": "geo error"}

            from . import seller as seller_mod

            evals = await asyncio.gather(*[_eval_guarded(c) for c in rep_cards.values()])
            for (canonical, card), ev in zip(rep_cards.items(), evals):
                card.extra.update({
                    "work_mode": ev["work_mode"],
                    "distance_miles": ev.get("distance_miles"),
                    "location_status": ev.get("location_status"),
                    "location_reason": ev.get("reason"),
                })
            ev_by_canonical = dict(zip(rep_cards.keys(), evals))
            for canonical, card in rep_cards.items():
                seller, s_reason = seller_mod.classify(
                    card.company, card.description,
                    card.extra.get("recruiter"))
                card.extra["posted_by"] = seller
                card.extra["posted_by_reason"] = s_reason

            hashmap = await db.upsert_jobs(conn, all_cards, resolution)
            await db.ensure_seed_blocklist(conn, settings)
            blocked_companies, blocked_keywords = await db.load_blocked_norms(conn)
            from . import training as training_mod
            from . import requirements as req_mod

            # collect the best CV text for requirements matching
            cur = await conn.execute(
                "SELECT full_text FROM cvs ORDER BY length(full_text) DESC LIMIT 3")
            cv_rows = await cur.fetchall()
            cv_text = "\n".join((r["full_text"] or "") for r in cv_rows)

            best: dict[str, float] = {}
            job_ids: dict[str, str] = {}
            blocked_skipped = 0
            training_skipped = 0
            location_skipped = 0
            req_skipped = 0
            for card in all_cards:
                ch = job_hash(card.title, card.company, card.location_text, card.salary_min)
                canonical = resolution.get(ch, ch)
                hit = db.is_blocked(card.company,
                                    f"{card.title} {card.description or ''}",
                                    blocked_companies, blocked_keywords)
                if hit:
                    blocked_skipped += 1
                    job_ids.pop(canonical, None)
                    best.pop(canonical, None)
                    continue
                is_training, t_reason = training_mod.classify(card)
                if is_training:
                    card.extra["training_offer"] = True
                    card.extra["training_reason"] = t_reason
                    if params.exclude_training:
                        training_skipped += 1
                        job_ids.pop(canonical, None)
                        best.pop(canonical, None)
                        continue
                # work-mode-aware mismatch drop (the user's rule: a daily-on-site
                # job beyond the commute is not worth showing)
                ev = ev_by_canonical.get(canonical, {})
                wm = ev.get("work_mode") or card.extra.get("work_mode")
                ls = ev.get("location_status")
                if params.drop_mismatch and ls == "mismatch" and wm != "remote":
                    location_skipped += 1
                    job_ids.pop(canonical, None)
                    best.pop(canonical, None)
                    continue
                # hard-requirements check against the CVs
                if card.description and len(card.description) > 100:
                    reqs = req_mod.extract_requirements(card.description)
                    match = req_mod.match_requirements(reqs, cv_text)
                    card.extra["requirements_status"] = match["status"]
                    card.extra["requirements_unmet"] = [
                        r["value"] for r in match["unmet"]]
                    if match["status"] == "gap" and params.drop_req_gap:
                        req_skipped += 1
                        job_ids.pop(canonical, None)
                        best.pop(canonical, None)
                        continue
                job_ids[canonical] = hashmap[ch][0]
                score, _ = score_job(card, params.query, params.min_salary,
                                     params.work_from_home,
                                     ev_by_canonical[canonical].get("location_status"))
                if canonical not in best or score > best[canonical]:
                    best[canonical] = score
            search_id = db.new_id("sch")
            await db.create_search(conn, search_id, params, {
                "providers": {r.provider: r.model_dump() for r in reports},
                "duplicates_merged": merged,
                "blocked_skipped": blocked_skipped,
                "training_skipped": training_skipped,
                "location_skipped": location_skipped,
                "req_skipped": req_skipped,
                "home": home,
                "response_policy": response_policy,
            })
            rows = sorted(((job_ids[c], s) for c, s in best.items()), key=lambda x: -x[1])
            await db.add_search_results(conn, search_id, [
                (job_id, score, rank) for rank, (job_id, score) in enumerate(rows, 1)
            ])
            await conn.commit()
        return await _page_results(search_id, response_policy=response_policy, offset=0,
                                   reports=reports, merged=merged,
                                   blocked_skipped=blocked_skipped,
                                   training_skipped=training_skipped,
                                   location_skipped=location_skipped)

    async def _saved_response_policy(search_id: str) -> dict[str, Any] | None:
        """Load the page policy saved with a previous search."""
        async with db.connect(settings.db_path) as conn:
            cur = await conn.execute("SELECT stats FROM searches WHERE id=?", (search_id,))
            row = await cur.fetchone()
        if not row or not row["stats"]:
            return None
        try:
            policy = json.loads(row["stats"]).get("response_policy")
        except (TypeError, ValueError):
            return None
        if not isinstance(policy, dict) or policy.get("profile") not in {
            "compact", "balanced", "wide"
        }:
            return None
        # Re-resolve instead of trusting old or manually edited numeric bounds.
        return _resolve_response_policy(
            policy["profile"], policy.get("context_window"), policy.get("page_limit")
        )

    async def _page_results(search_id: str, response_policy: dict[str, Any], offset: int = 0,
                            reports: list | None = None, merged: int | None = None,
                            blocked_skipped: int | None = None,
                            training_skipped: int | None = None,
                            location_skipped: int | None = None) -> dict:
        limit = response_policy["page_limit"]
        async with db.connect(settings.db_path) as conn:
            rows, total = await db.get_search_results(conn, search_id, limit, offset)
            if not rows and offset == 0:
                return {"error": "search produced 0 results — broaden the query "
                        "or check provider reports"}
            if not rows:
                return {"search_id": search_id, "total": total, "results": [],
                        "note": "end of results"}
            src_map = await db.sources_for_hashes(conn, [r["content_hash"] for r in rows])
            briefs = []
            for r in rows:
                b = db.brief_from_row(r, r["rank"], False,
                                      [s["source"] for s in src_map.get(r["content_hash"], [])])
                method, _u, _c = tracker_mod._apply_method(r)
                b["apply_method"] = method
                briefs.append(b)
            out = {
                "candidate": settings.candidate_name,
                "search_id": search_id, "total": total,
                "showing": f"{offset + 1}-{offset + len(rows)} of {total}",
                "next_offset": offset + limit if offset + limit < total else None,
                "response_policy": response_policy,
                "results": briefs,
                "hints": [
                    "give job_ids to the user; apply via start_application",
                    "when presenting the list ALWAYS show posted_by "
                    "(agency vs direct employer) for every job",
                    "for each job show a SHORT DESCRIPTION (2-3 sentences "
                    "from the description field; if null, call "
                    "fetch_job_description for the top picks first)",
                    "requirements_status=gap → the user does NOT meet a "
                    "hard requirement — do not propose the job",
                    "location_status=mismatch → too far & not remote, do not apply "
                    "without user approval",
                ],
            }
            if out["next_offset"] is not None:
                out["next_page"] = {
                    "search_id": search_id,
                    "offset": out["next_offset"],
                    "response_profile": response_policy["profile"],
                    "limit": limit,
                }
            if reports is not None:
                out["provider_reports"] = [r.model_dump() for r in reports]
            if merged is not None:
                out["duplicates_merged"] = merged
            if blocked_skipped:
                out["blocked_skipped"] = blocked_skipped
            if training_skipped:
                out["training_offers_skipped"] = training_skipped
                out["note_training"] = ("paid training/course ads were excluded "
                                        "(pass include_training=true to see them)")
            if location_skipped:
                out["location_skipped"] = location_skipped
                out["note_location"] = (
                    "on-site jobs beyond the active profile's commute limit were dropped "
                    f"(work-mode-aware: daily ≤{settings.daily_commute_miles}mi, "
                    f"hybrid/field ≤{settings.occasional_commute_miles}mi, remote unlimited)"
                )
            return out

    @mcp.tool()
    async def search_jobs(
        query: str | None = None,
        profile: str | None = None,
        search_id: str | None = None,
        offset: int = 0,
        limit: int | None = None,
        response_profile: ResponseProfile = "auto",
        context_window: int | None = None,
        location: str | None = None,
        radius_miles: int | None = None,
        max_days_old: int | None = None,
        work_from_home: bool | None = None,
        min_salary: int | None = None,
        sources: list[str] | None = None,
        limit_per_source: int | None = None,
        location_policy: str | None = None,
        include_training: bool = False,
        drop_mismatch: bool | None = None,
        drop_req_gap: bool | None = None,
        enrich_descriptions: int = 0,
    ) -> dict:
        """Run a UK job search OR page an earlier one. Fresh search: pass
        query ('Data Analyst'…) or profile ('data_analytics'/'field_geologist').
        Paging: pass search_id (+offset from next_offset), preferably by copying
        the returned next_page object. Response sizing: pass context_window and
        auto selects compact (<=80k), balanced (80k-300k), or wide (>=300k).
        Local Qwen 3.8 27B uses context_window=78000. You may instead pass
        response_profile explicitly. Optional limit requests a page size within
        that profile's safety cap (8/20/50). With no sizing signal the default is
        balanced (12). All remaining results stay stored. Results are ranked,
        cross-board duplicates merged (sources[]), with memory flags
        (already_applied, application_status) and location intelligence
        (work_mode, distance_miles from the active profile's home, location_status
        ok|mismatch|caution|unknown). PAID TRAINING/COURSE ADS (where you pay
        them, e.g. Netcom-style 'trainee' course marketing) are excluded
        automatically — training_offers_skipped shows how many; set
        include_training=true only if the user explicitly wants courses.
        location_policy: 'auto' (default), 'uk_wide', 'commute_only'.
        work_from_home=true restricts to remote.
        """
        if search_id:
            if response_profile == "auto" and context_window is None and limit is None:
                response_policy = await _saved_response_policy(search_id)
            else:
                response_policy = _resolve_response_policy(
                    response_profile, context_window, limit
                )
            response_policy = response_policy or _resolve_response_policy()
            return await _page_results(search_id, response_policy, offset)
        response_policy = _resolve_response_policy(response_profile, context_window, limit)
        params_kwargs: dict[str, Any] = {}
        if profile:
            prof = settings.search_profiles.get(profile)
            if not prof:
                return {"error": f"unknown profile '{profile}' (available: "
                        f"{list(settings.search_profiles)})"}
            params_kwargs.update(prof)
        if query:
            params_kwargs["query"] = query
        if location is not None:
            params_kwargs["location"] = location
        params_kwargs.setdefault("location", settings.default_location)
        params_kwargs.setdefault("radius_miles",
                                 radius_miles or settings.default_radius_miles)
        params_kwargs.setdefault("max_days_old",
                                 max_days_old or settings.default_max_days_old)
        params_kwargs["work_from_home"] = (work_from_home
                                           if work_from_home is not None
                                           else settings.default_work_from_home)
        if min_salary:
            params_kwargs["min_salary"] = min_salary
        if sources:
            params_kwargs["sources"] = sources
        params_kwargs["limit_per_source"] = (limit_per_source
                                             or settings.default_limit_per_source)
        if location_policy:
            params_kwargs["location_policy"] = location_policy
        params_kwargs["exclude_training"] = (not include_training
                                             and settings.exclude_training)
        if drop_mismatch is not None:
            params_kwargs["drop_mismatch"] = drop_mismatch
        if drop_req_gap is not None:
            params_kwargs["drop_req_gap"] = drop_req_gap
        if enrich_descriptions:
            params_kwargs["_enrich"] = enrich_descriptions
        params_kwargs.pop("radius", None)
        params = SearchParams(**{k: v for k, v in params_kwargs.items() if v is not None})
        if not params.query:
            return {"error": "query or profile is required"}
        return await _run_search(params, response_policy)

    @mcp.tool()
    async def get_job(
        job_ids: list[str],
        include_description: bool = False,
        response_profile: ResponseProfile = "auto",
        context_window: int | None = None,
    ) -> dict:
        """Details for selected jobs with adaptive batch sizing. Pass the same
        context_window/response_profile used for search_jobs. Compact returns up
        to 5 summaries or 2 full descriptions; balanced 12/5; wide 30/12.
        Request full descriptions only for plausible finalists. No job is deleted
        when a supplied job_ids list is truncated; request the remainder later."""
        out_jobs = []
        response_policy = _resolve_response_policy(
            response_profile, context_window, None
        )
        caps = {
            "compact": {False: 5, True: 2},
            "balanced": {False: 12, True: 5},
            "wide": {False: 30, True: 12},
        }
        max_jobs = caps[response_policy["profile"]][include_description]
        async with db.connect(settings.db_path) as conn:
            for jid in job_ids[:max_jobs]:
                job = await db.get_job(conn, jid)
                if not job:
                    out_jobs.append({"job_id": jid, "error": "unknown"})
                    continue
                srcs = await db.job_sources(conn, job["content_hash"])
                method, apply_url, cautions = tracker_mod._apply_method(job)
                try:
                    extra = json.loads(job.pop("extra") or "{}")
                except (TypeError, ValueError):
                    extra = {}
                job.update({
                    "sources_detail": srcs,
                    "apply_method": method,
                    "apply_url_resolved": apply_url,
                    "apply_cautions": cautions,
                    "site_playbook": tracker_mod.PLAYBOOKS.get(
                        method, tracker_mod.PLAYBOOKS["website_form"]),
                    "work_mode": extra.get("work_mode"),
                    "distance_miles": extra.get("distance_miles"),
                    "location_status": extra.get("location_status"),
                    "location_reason": extra.get("location_reason"),
                    "posted_by": extra.get("posted_by"),
                    "posted_by_reason": extra.get("posted_by_reason"),
                    "training_offer": extra.get("training_offer", False),
                    "requirements_status": extra.get("requirements_status"),
                    "requirements_unmet": extra.get("requirements_unmet"),
                })
                if not include_description:
                    job["description"] = (job.get("description") or "")[:400] + "…"
                app = await db.application_for_job(conn, jid)
                if app:
                    job["application"] = {k: app[k] for k in
                                          ("id", "status", "submitted_at", "updated_at")}
                out_jobs.append(job)
        metadata = {
            "profile": response_policy["profile"],
            "context_window": context_window,
            "max_jobs": max_jobs,
            "requested_jobs": len(job_ids),
            "returned_jobs": len(out_jobs),
            "remaining_job_ids": job_ids[max_jobs:],
        }
        if len(out_jobs) > 1:
            return {"response_policy": metadata, "jobs": out_jobs}
        if not out_jobs:
            return {"response_policy": metadata, "jobs": []}
        out_jobs[0]["response_policy"] = metadata
        return out_jobs[0]

    @mcp.tool()
    async def manage_blocklist(action: str = "list", kind: str = "company",
                               value: str | None = None,
                               reason: str | None = None) -> dict:
        """Employer/recruiter exclusion memory. When the user says 'never apply
        to <company>' → action=add, kind=company, value=<name>, reason=<quote>.
        kind=keyword blocks any vacancy whose title/description contains it.
        action: add | remove | list. Blocked jobs are hidden from search
        results and start_application refuses them until removed."""
        async with db.connect(settings.db_path) as conn:
            if action == "add" and value:
                added = await db.blocklist_add(conn, kind, value, reason)
                await conn.commit()
                return {"ok": True, "added": added,
                        "note": "" if added else "already blocklisted",
                        "blocklist": await db.blocklist_list(conn)}
            if action == "remove" and value:
                removed = await db.blocklist_remove(conn, kind, value)
                await conn.commit()
                return {"ok": removed, "removed": removed,
                        "blocklist": await db.blocklist_list(conn)}
            return {"blocklist": await db.blocklist_list(conn)}

    @mcp.tool()
    async def fetch_job_description(job_ids: list[str]) -> dict:
        """Fetch FULL job descriptions by opening pages in the browser.
        Accepts 1-10 job_ids (BATCH - prefer batching to save round-trips).
        Use for jobs where the parser returned null/short descriptions
        (common on Totaljobs/Reed). Stores descriptions back and re-runs the
        requirements check against your CVs. Also detects closed vacancies
        ("no longer accepting applications")."""
        from . import requirements as req_mod
        from .browser import BrowserError, get_session

        if not job_ids:
            return {"error": "job_ids list is empty"}
        job_ids = [j for j in job_ids if j][:10]
        if not job_ids:
            return {"error": "job_ids contained no valid ids"}
        results = []
        jobs = []
        async with db.connect(settings.db_path) as conn:
            for jid in job_ids:
                job = await db.get_job(conn, jid)
                if not job:
                    results.append({"job_id": jid, "error": "unknown"})
                else:
                    jobs.append(job)
        if not jobs:
            return {"results": results}
        sess = get_session(settings)
        for job in jobs:
            jid = job["id"]
            url = job.get("apply_url") or job.get("url")
            if not url:
                results.append({"job_id": jid, "error": "no URL"})
                continue
            try:
                await asyncio.wait_for(sess.open(url), timeout=25)
                await sess._active().wait_for_timeout(1500)
                result = await sess.evaluate(
                    "() => (document.body.innerText || '').replace(/\\s+/g, ' ')")
                full_text = (result.get("result") or "")
                desc_start = full_text.find("Job description")
                if desc_start == -1:
                    desc_start = full_text.find("About the role")
                if desc_start == -1:
                    desc_start = 500
                desc_end = full_text.rfind("Apply")
                if desc_end == -1 or desc_end <= desc_start:
                    desc_end = len(full_text)
                description = full_text[desc_start:desc_end].strip()[:8000]
                closed = any(marker in full_text.lower() for marker in
                             ("no longer accepting", "can no longer apply",
                              "vacancy has been closed", "position has been filled",
                              "closing date has passed"))
            except (BrowserError, asyncio.TimeoutError, Exception) as exc:
                results.append({"job_id": jid,
                                "error": f"{type(exc).__name__}: {str(exc)[:120]}"})
                continue

            async with db.connect(settings.db_path) as conn:
                await conn.execute(
                    "UPDATE jobs SET description=? WHERE id=?", (description, jid))
                cur = await conn.execute(
                    "SELECT full_text FROM cvs ORDER BY length(full_text) DESC LIMIT 3")
                cv_rows = await cur.fetchall()
                cv_text = "\n".join((r["full_text"] or "") for r in cv_rows)
                reqs = req_mod.extract_requirements(description)
                match = req_mod.match_requirements(reqs, cv_text)
                cur2 = await conn.execute("SELECT extra FROM jobs WHERE id=?", (jid,))
                row = await cur2.fetchone()
                extra = {}
                try:
                    extra = json.loads(row["extra"] or "{}")
                except (TypeError, ValueError):
                    pass
                extra["requirements_status"] = match["status"]
                extra["requirements_unmet"] = [r["value"] for r in match["unmet"]]
                await conn.execute("UPDATE jobs SET extra=? WHERE id=?",
                                   (json.dumps(extra, ensure_ascii=False), jid))
                await conn.commit()
            results.append({
                "job_id": jid,
                "description_length": len(description),
                "description_preview": description[:400],
                "requirements_status": match["status"],
                "requirements_unmet": [r["value"] for r in match["unmet"]],
                "vacancy_closed": closed or None,
            })
        return {"results": results}

    @mcp.tool()
    async def list_stored_jobs(
        query: str | None = None,
        company: str | None = None,
        location: str | None = None,
        source: str | None = None,
        days_old: int = 7,
        limit: int = 20,
    ) -> dict:
        """Search the LOCAL job database by criteria (no board calls). Use
        INSTEAD of Bash/sqlite: find all jobs from a company, filter by title
        keyword, location, source board, or freshness. Returns job_ids +
        titles + requirements status - then get_job / fetch_job_description
        for detail."""
        async with db.connect(settings.db_path) as conn:
            sql = ("SELECT id, title, company, location_text, salary_raw, "
                   "posted_at, source, extra, description FROM jobs "
                   "WHERE last_seen >= datetime('now', ?)")
            args: list = [f"-{days_old} days"]
            if query:
                sql += " AND (title LIKE ? OR description LIKE ?)"
                args += [f"%{query}%", f"%{query}%"]
            if company:
                sql += " AND company LIKE ?"
                args.append(f"%{company}%")
            if location:
                sql += " AND location_text LIKE ?"
                args.append(f"%{location}%")
            if source:
                sql += " AND source = ?"
                args.append(source)
            sql += " ORDER BY posted_at DESC LIMIT ?"
            args.append(limit)
            cur = await conn.execute(sql, args)
            rows = []
            for r in await cur.fetchall():
                extra = {}
                try:
                    extra = json.loads(r["extra"] or "{}")
                except (TypeError, ValueError):
                    pass
                rows.append({
                    "job_id": r["id"], "title": r["title"],
                    "company": r["company"], "location": r["location_text"],
                    "salary": r["salary_raw"], "posted_at": r["posted_at"],
                    "source": r["source"],
                    "requirements_status": extra.get("requirements_status"),
                    "requirements_unmet": extra.get("requirements_unmet"),
                    "has_description": bool(r["description"]
                                            and len(r["description"]) > 200),
                })
            return {"jobs": rows, "count": len(rows),
                    "note": "get_job for full records; fetch_job_description "
                            "to enrich missing ones"}

    @mcp.tool()
    async def submit_job_observations(
        observations: list[dict],
        observation_type: str = "search_cards",
        task_id: str | None = None,
        search_id: str | None = None,
    ) -> dict:
        """Feed jobs found via the harness's own browser (Indeed, CV-Library,
        LinkedIn, Glassdoor) into the same dedup + ranking store. Each
        observation: {source, url, title, company, location_text, salary_raw,
        description, posted_at?}. Returns stored count + duplicates merged."""
        from .domain import ObservationIn
        from .textutils import parse_salary

        cards = []
        for o in observations:
            try:
                obs = ObservationIn(**o)
            except Exception as exc:  # noqa: BLE001
                return {"error": f"invalid observation {o.get('title')!r}: {exc}"}
            sal = parse_salary(obs.salary_raw)
            cards.append(JobCard(
                source=obs.source, source_job_id=obs.source_job_id, url=obs.url,
                apply_url=obs.url, title=obs.title, company=obs.company,
                location_text=obs.location_text, salary_raw=obs.salary_raw,
                salary_min=sal[0], salary_max=sal[1], salary_period=sal[2],
                contract_type=obs.contract_type, work_from_home=obs.work_from_home,
                description=obs.description, posted_at=obs.posted_at,
                extra={"observed": True},
            ))
        async with db.connect(settings.db_path) as conn:
            pool = await dedup.load_pool(conn)
            resolution, merged = dedup.resolution_map(cards, pool)
            await db.upsert_jobs(conn, cards, resolution)
            await db.log_observation(conn, task_id, observation_type,
                                     observations[0].get("source") if observations else None,
                                     len(cards), search_id)
            await conn.commit()
        return {"stored": len(cards), "duplicates_merged": merged,
                "note": "rank them together with HTTP sources via search_jobs; "
                        "details via get_job"}

    # ----------------------------------------------------------------- cv ----
    @mcp.tool()
    async def list_cvs(job_id: str | None = None) -> dict:
        """CVs indexed from the active profile's local folder.
        Pass job_id to get per-CV recommendation scores for that job."""
        async with db.connect(settings.db_path) as conn:
            cvs = await db.list_cvs(conn)
            recs = []
            if job_id:
                job = await db.get_job(conn, job_id)
                if not job:
                    return {"error": f"unknown job_id {job_id}"}
                recs = await _recommend_cv(conn, job, limit=5)
            return {
                "candidate": settings.candidate_name,
                "cv_dir": str(settings.cv_dir),
                "cvs": cvs,
                "recommendations_for_job": recs or None,
                "note": "Copy CV files into cv_dir manually; call sync_cvs after changes",
            }

    @mcp.tool()
    async def sync_cvs(force: bool = False) -> dict:
        """Index CV files manually copied into the active profile's cv_dir.
        Unchanged files are skipped unless force=true. There is no cloud sync."""
        result = await _index_cvs(settings, force=force)
        return {
            "candidate": settings.candidate_name,
            "cv_dir": str(settings.cv_dir),
            "storage": "local_manual",
            **result,
        }

    # ------------------------------------------------------------- apply ----
    @mcp.tool()
    async def start_application(job_id: str, cv_id: str | None = None,
                                notes: str | None = None) -> dict:
        """Begin applying to a job. Returns the full plan: URL, apply method +
        site playbook, chosen CV (+alternatives), applicant profile values,
        step list, cautions (location mismatch, LinkedIn ToS…). REFUSES to
        create a second application for a job that already has one — that is
        the anti-double-apply memory (works across boards via dedup)."""
        async with db.connect(settings.db_path) as conn:
            result = await tracker_mod.start_application(conn, settings, job_id, cv_id, notes)
            await conn.commit()
        result["candidate"] = settings.candidate_name
        return result

    @mcp.tool()
    async def record_application(
        application_id: str,
        status: str,
        notes: str | None = None,
        evidence: dict | None = None,
        cover_letter: str | None = None,
    ) -> dict:
        """Update an application: status planned|applying|submitted|interview|
        offer|rejected|withdrawn|failed; notes; evidence {screenshot: path,…};
        cover_letter text."""
        async with db.connect(settings.db_path) as conn:
            try:
                app = await db.update_application(conn, application_id, status, notes,
                                                  cover_letter, evidence)
            except ValueError as exc:
                return {"error": str(exc)}
            if app is None:
                return {"error": f"unknown application_id {application_id}"}
            await conn.commit()
            return {"ok": True, "application": app}

    @mcp.tool()
    async def list_applications(status: str | None = None, limit: int = 50) -> dict:
        """Application history (job, company, status, dates) — the long-term
        memory of everything already applied to."""
        async with db.connect(settings.db_path) as conn:
            apps = await db.list_applications(conn, status, limit)
            for a in apps:
                a.pop("evidence", None)
                a.pop("cover_letter", None)
            return {"applications": apps, "statuses": sorted(db.APP_STATUSES)}

    @mcp.tool()
    async def check_applied(
        job_id: str | None = None, url: str | None = None,
        title: str | None = None, company: str | None = None,
    ) -> dict:
        """'Have we already applied here?' — by job_id, exact URL, or fuzzy
        title+company (catches the same vacancy re-found on another board).
        ALWAYS call before a manual browser application."""
        async with db.connect(settings.db_path) as conn:
            if job_id:
                app = await db.application_for_job(conn, job_id)
                matches = [app] if app else []
            else:
                matches = await db.find_applications_like(conn, url, title, company)
                for m in matches:
                    m.pop("evidence", None)
            return {
                "matches": matches,
                "verdict": ("ALREADY APPLIED — do not submit again"
                            if any(m.get("submitted_at") for m in matches)
                            else ("in progress (planned/applying)" if matches
                                  else "no prior application")),
            }

    @mcp.tool()
    async def make_cover_letter(text: str, name: str = "Cover_Letter") -> dict:
        """Write a cover letter as a DOCX file into CV_collection and return
        its path — ready for browser_upload as a supporting file. Boards like
        Totaljobs reject files smaller than 8KB, so the document is padded
        through metadata if the text alone is too short. Prefer DOCX over PDF
        here (PDFs from minimal text are usually under the limit)."""
        import re as _re

        from docx import Document

        safe = _re.sub(r"[^A-Za-z0-9_-]+", "_", name)[:50] or "Cover_Letter"
        out = settings.cv_dir / f"{safe}.docx"
        settings.cv_dir.mkdir(parents=True, exist_ok=True)

        def _build():
            doc = Document()
            first = True
            for ln in text.splitlines():
                if ln.strip() and first:
                    doc.add_heading(ln.strip(), level=0)
                    first = False
                elif ln.strip():
                    doc.add_paragraph(ln.strip())
            doc.save(str(out))
            # Totaljobs minimum: supporting files must be > 8KB. DOCX
            # properties cap at 255 chars, so pad with empty paragraphs.
            while out.stat().st_size < 9_000:
                doc.add_paragraph("")
                doc.save(str(out))
            return out.stat().st_size

        size = await asyncio.to_thread(_build)
        return {"path": str(out), "size_bytes": size,
                "min_board_limit": 8192,
                "note": "upload with browser_upload as a supporting file"}

    # ---------------------------------------------------------- browser ----
    @mcp.tool()
    async def browser_login(url: str) -> dict:
        """Ensure we're signed in on a job board before applying. Opens the
        site; if signed out, walks 'Continue with Google' and picks the
        pre-approved account (active profile auth.google_account — the user allows
        this WITHOUT asking). Returns logged_in; needs_user=true on 2FA/
        captcha/consent — then stop and ask the user to finish in the window."""
        account = settings.auth.get("google_account") \
            if settings.auth.get("auto_google_signin", True) else None
        try:
            return await get_session(settings).login_flow(url, account)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_open(url: str, headless: bool | None = None) -> dict:
        """Open a URL in the persistent application browser; returns the first
        snapshot. Logins survive between runs. Headed by default so the user
        can handle 2FA/captcha."""
        try:
            return await get_session(settings).open(url, headless)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_snapshot(focus: str | None = None,
                               filter_text: str | None = None,
                               text_chars: int = 800,
                               modal_only: bool = False) -> dict:
        """Look at the page: numbered interactive elements + text. focus:
        'inputs'|'buttons'|'links'; filter_text narrows by name (find 'Apply');
        text_chars=0 → elements only, 6000 → reading mode. modal_only=true →
        ONLY the active dialog/wizard (question text + controls, hidden
        templates excluded) — use for apply wizards, then browser_set/click
        by number."""
        try:
            return await get_session(settings).snapshot(focus, filter_text, text_chars, modal_only)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_form() -> dict:
        """The dominant form with human labels per field (n, tag, type, label,
        required, options) + likely submit buttons. Use once per application
        page, then browser_set fields by number."""
        try:
            return await get_session(settings).form()
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_click(n: int, timeout_ms: int | None = None) -> dict:
        """Click element #n → returns the FRESH snapshot (popups become the
        active tab automatically)."""
        try:
            return await get_session(settings).click(n, timeout_ms)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_set(n: int, value: str | list[str] | bool) -> dict:
        """Set element #n to a value — one verb for every control: inputs →
        fill; selects → pick option value; checkboxes/radios → true/false →
        returns fresh snapshot."""
        try:
            return await get_session(settings).set(n, value)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_type(n: int, text: str, submit: bool = False) -> dict:
        """Type into #n key-by-key (fires JS handlers); submit=true presses
        Enter → fresh snapshot. Use for search boxes and chat fields."""
        try:
            return await get_session(settings).type_text(n, text, submit)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_upload(n: int, file_path: str) -> dict:
        """Click upload control #n and send a local file (the CV path from
        start_application) → fresh snapshot."""
        try:
            return await get_session(settings).upload(n, file_path)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_press(key: str) -> dict:
        """Press a page-level key (Enter, Tab, Escape) → fresh snapshot."""
        try:
            return await get_session(settings).press(key)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_wait(seconds: float | None = None, text: str | None = None,
                           text_gone: str | None = None) -> dict:
        """Wait for time / text to appear ('Application submitted') / text to
        disappear (spinners) → returns a snapshot."""
        try:
            return await get_session(settings).wait(seconds, text, text_gone)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_screenshot(name: str | None = None,
                                 full_page: bool = False) -> dict:
        """Save a PNG to data/screenshots — pass the path as evidence to
        record_application."""
        try:
            return await get_session(settings).screenshot(name, full_page)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_eval(js: str) -> dict:
        """Evaluate JS on the page (extract hidden JSON, scroll, dismiss cookie
        banners)."""
        try:
            return await get_session(settings).evaluate(js)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_tabs(action: str = "list", index: int | None = None) -> dict:
        """Tabs: list / select / close. Click-opened popups become active
        automatically — use this to switch back."""
        try:
            return await get_session(settings).tabs(action, index)
        except Exception as exc:  # noqa: BLE001 - browser errors must not kill the tool call
            return {"error": str(exc)}

    @mcp.tool()
    async def browser_close() -> dict:
        """Close the browser (the login profile persists on disk)."""
        return await get_session(settings).close()


async def run_stdio(settings: Settings | None = None) -> None:
    mcp, _ = create_server(settings)
    await mcp.run_stdio_async()
